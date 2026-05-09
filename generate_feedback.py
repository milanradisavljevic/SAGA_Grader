#!/usr/bin/env python3
"""
Korrektur-Feedback Generator

Erzeugt Feedback-DOCX aus fertigen Analyse-JSONs in /output/feedback_data.
"""

from __future__ import annotations

import argparse
import json
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# Farbpalette
C_PRIMARY = RGBColor(0x1F, 0x49, 0x7D)  # Dunkelblau  – Header, Standard
C_HEADER = RGBColor(0x2E, 0x4D, 0x8A)  # Dunkelblau  – Abschnittstitel
C_DIVIDER = RGBColor(0x99, 0x99, 0x99)  # Grau        – Trennlinien
C_STRENGTH = RGBColor(0x00, 0x70, 0x00)  # Dunkelgrün  – Stärken
C_WEAKNESS = RGBColor(0xC0, 0x00, 0x00)  # Dunkelrot   – Schwächen
C_SUGGESTION = RGBColor(0x00, 0x46, 0x7F)  # Blau        – Verbesserungsvorschläge
C_GRADE_1 = RGBColor(0x70, 0xC0, 0x70)  # Hellgrün   – Note 1
C_GRADE_2 = RGBColor(0x00, 0x80, 0x00)  # Dunkelgrün – Note 2
C_GRADE_3 = RGBColor(0xC0, 0xA0, 0x00)  # Gelb       – Note 3
C_GRADE_4 = RGBColor(0xE0, 0x60, 0x00)  # Orange     – Note 4
C_GRADE_5 = RGBColor(0xC0, 0x00, 0x00)  # Rot        – Note 5

# Rückwärtskompatible Aliasse (werden in Tests referenziert)
C_GRADE_GOOD = C_GRADE_2
C_GRADE_OK = C_GRADE_3
C_GRADE_FAIL = C_GRADE_5

_GRADE_COLORS = {1: C_GRADE_1, 2: C_GRADE_2, 3: C_GRADE_3, 4: C_GRADE_4, 5: C_GRADE_5}


def _note_color(note: int) -> RGBColor:
    """Notenfarbe: 1=hellgrün, 2=dunkelgrün, 3=gelb, 4=orange, 5=rot."""
    return _GRADE_COLORS.get(note, C_GRADE_5)


def _set_cell_bg(cell, hex_color: str) -> None:
    """Setzt die Hintergrundfarbe einer Tabellenzelle via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


GERMAN_ORDER = ["inhalt", "textstruktur", "ausdruck", "sprachrichtigkeit"]
ENGLISH_ORDER = [
    "task_achievement",
    "organisation_layout",
    "lexical_range_accuracy",
    "grammatical_range_accuracy",
]

ALIASES = {
    "stil_ausdruck": "ausdruck",
    "normative_sprachrichtigkeit": "sprachrichtigkeit",
    "erfuellung_aufgabenstellung": "task_achievement",
    "aufbau_layout": "organisation_layout",
    "wortschatz": "lexical_range_accuracy",
    "grammatik": "grammatical_range_accuracy",
}


@dataclass(slots=True)
class CriterionFeedback:
    """Bewertungsdaten fuer ein einzelnes Kriterium."""

    key: str
    stufe: str
    punkte: float
    staerken: list[str]
    schwaechen: list[str]
    vorschlaege: list[str]
    gewicht: float | None = None
    fehler_detail: list[str] | None = None
    fehlerschwerpunkte: list[str] | None = None
    rhetorische_figuren: list[str] | None = None


@dataclass(slots=True)
class GradeRecommendation:
    """Zusammenfassung der Notenempfehlung."""

    durchschnitt: float
    note: int
    bezeichnung: str
    begruendung: str


@dataclass(slots=True)
class FeedbackData:
    """Vollstaendige Daten fuer ein Feedback-Dokument."""

    datei: str
    schueler: str | None
    klasse: str | None
    textsorte: str
    fach: str
    schulstufe: str
    rubrik: str
    bewertung: list[CriterionFeedback]
    notenempfehlung: GradeRecommendation
    hinweise: list[str]


@dataclass(slots=True)
class ProjectPaths:
    """Projektpfade relativ zum Skriptstandort."""

    root: Path
    input_dir: Path
    output_dir: Path
    feedback_data_dir: Path
    fehlerlog: Path


def project_paths() -> ProjectPaths:
    """Ermittelt die Projektpfade relativ zum Skript."""

    root = Path(__file__).resolve().parent
    output_dir = root / "output"
    return ProjectPaths(
        root=root,
        input_dir=root / "input",
        output_dir=output_dir,
        feedback_data_dir=output_dir / "feedback_data",
        fehlerlog=output_dir / "fehlerlog.txt",
    )


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Fuegt eine Ueberschrift ein."""

    doc.add_heading(text, level=level)


def add_divider(doc: Document) -> None:
    """Fuegt einen dezenten Trenner ein."""

    paragraph = doc.add_paragraph("-" * 56)
    if paragraph.runs:
        paragraph.runs[0].font.color.rgb = C_DIVIDER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)


def add_section_header(doc: Document, text: str, color: RGBColor = C_HEADER) -> None:
    """Fuegt einen Abschnittstitel im bestaenden Stil ein."""

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = color
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(2)


def add_label(doc: Document, label: str, value: str) -> None:
    """Fuegt eine Label-Wert-Zeile ein."""

    paragraph = doc.add_paragraph()
    key_run = paragraph.add_run(f"{label}: ")
    key_run.bold = True
    key_run.font.size = Pt(10)
    value_run = paragraph.add_run(value)
    value_run.font.size = Pt(10)
    paragraph.paragraph_format.space_after = Pt(1)


def add_bullet(
    doc: Document, text: str, indent: int = 1, color: RGBColor | None = None
) -> None:
    """Fuegt einen Listenpunkt ein."""

    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    if color is not None:
        run.font.color.rgb = color
    paragraph.paragraph_format.left_indent = Inches(0.3 * indent)
    paragraph.paragraph_format.space_after = Pt(1)


def add_body(doc: Document, text: str) -> None:
    """Fuegt einen normalen Absatz ein."""

    paragraph = doc.add_paragraph(text)
    for run in paragraph.runs:
        run.font.size = Pt(10)
    paragraph.paragraph_format.space_after = Pt(2)


def setup_page(doc: Document, config: dict[str, Any] | None = None) -> None:
    """Setzt A4-Seitenformat, Raender, Fusszeile und optionale Kopfzeile."""

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = fp.add_run(
        f"Erstellt mit SAGA — {date.today().strftime('%d.%m.%Y')}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = C_DIVIDER

    if config:
        docx_cfg = config.get("docx", {})
        teacher = docx_cfg.get("teacher_name", "")
        school = docx_cfg.get("school_name", "")
        if teacher or school:
            header = section.header
            header.is_linked_to_previous = False
            parts = [p for p in [teacher, school] if p]
            hp = header.paragraphs[0]
            header_run = hp.add_run(" | ".join(parts))
            header_run.font.size = Pt(8)
            header_run.font.color.rgb = C_DIVIDER


def add_document_header(
    doc: Document, data: FeedbackData, config: dict[str, Any] | None = None
) -> None:
    """Fuegt den Metadaten-Header mit optionaler Metadaten-Tabelle ein."""

    docx_cfg = (config or {}).get("docx", {})

    logo_path = docx_cfg.get("logo_path", "")
    if logo_path:
        lp = Path(logo_path)
        if not lp.is_absolute():
            lp = project_paths().root / lp
        if lp.exists():
            doc.add_picture(str(lp), width=Cm(4))
            doc.add_paragraph()

    add_heading(doc, "KORREKTUR-FEEDBACK", level=1)

    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.autofit = True

    rows_data: list[tuple[str, str]] = []
    rows_data.append(("Datei", data.datei))
    if data.schueler:
        val = data.schueler if not data.klasse else f"{data.schueler}, {data.klasse}"
        rows_data.append(("Schueler/in", val))
    rows_data.append(("Fach", data.fach))
    rows_data.append(("Schulstufe", data.schulstufe))
    rows_data.append(("Textsorte", data.textsorte))
    rows_data.append(("Rubrik", data.rubrik))
    rows_data.append(("Datum", date.today().strftime("%d.%m.%Y")))

    teacher = docx_cfg.get("teacher_name", "")
    school = docx_cfg.get("school_name", "")
    if teacher:
        rows_data.append(("Lehrer/in", teacher))
    if school:
        rows_data.append(("Schule", school))

    for label, value in rows_data:
        row = meta_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    add_divider(doc)


def parse_args() -> argparse.Namespace:
    """Liest die CLI-Argumente ein."""

    parser = argparse.ArgumentParser(
        description="Erzeugt Feedback-DOCX aus JSON-Dateien."
    )
    parser.add_argument(
        "--file", help="Verarbeitet nur eine JSON-Datei aus /output/feedback_data."
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Ueberschreibt bestehende _feedback.docx-Dateien.",
    )
    parser.add_argument(
        "--dry-run", action="store_true", help="Zeigt nur, was verarbeitet wuerde."
    )
    return parser.parse_args()


def canonical_key(raw_key: str) -> str:
    """Normalisiert Kriterienschluessel."""

    return ALIASES.get(raw_key, raw_key)


def ensure_list(value: Any, field_name: str) -> list[str]:
    """Stellt sicher, dass ein Feld eine String-Liste ist."""

    if value is None:
        return []
    if not isinstance(value, list) or not all(isinstance(item, str) for item in value):
        raise ValueError(f"Feld '{field_name}' muss eine Liste aus Strings sein.")
    return value


def parse_criterion(key: str, payload: dict[str, Any]) -> CriterionFeedback:
    """Validiert ein einzelnes Kriterium."""

    required = {"stufe", "punkte", "staerken", "schwaechen", "vorschlaege"}
    missing = sorted(required - payload.keys())
    if missing:
        raise ValueError(f"Kriterium '{key}' fehlt: {', '.join(missing)}")

    stufe = payload["stufe"]
    if not isinstance(stufe, (str, int, float)):
        raise ValueError(f"Kriterium '{key}' hat eine ungueltige Stufe.")

    punkte = payload["punkte"]
    if not isinstance(punkte, (int, float)):
        raise ValueError(f"Kriterium '{key}' hat ungueltige Punkte.")

    gewicht = payload.get("gewicht")
    if gewicht is not None and not isinstance(gewicht, (int, float)):
        raise ValueError(f"Kriterium '{key}' hat ein ungueltiges Gewicht.")

    return CriterionFeedback(
        key=canonical_key(key),
        stufe=str(stufe),
        punkte=float(punkte),
        gewicht=float(gewicht) if gewicht is not None else None,
        staerken=ensure_list(payload["staerken"], f"{key}.staerken"),
        schwaechen=ensure_list(payload["schwaechen"], f"{key}.schwaechen"),
        vorschlaege=ensure_list(payload["vorschlaege"], f"{key}.vorschlaege"),
        fehler_detail=ensure_list(payload.get("fehler_detail"), f"{key}.fehler_detail")
        or None,
        fehlerschwerpunkte=ensure_list(
            payload.get("fehlerschwerpunkte"), f"{key}.fehlerschwerpunkte"
        )
        or None,
        rhetorische_figuren=ensure_list(
            payload.get("rhetorische_figuren"), f"{key}.rhetorische_figuren"
        )
        or None,
    )


def parse_feedback_data(payload: dict[str, Any]) -> FeedbackData:
    """Validiert die JSON-Struktur und wandelt sie in Dataklassen um."""

    required = {
        "datei",
        "textsorte",
        "fach",
        "schulstufe",
        "rubrik",
        "bewertung",
        "notenempfehlung",
    }
    missing = sorted(required - payload.keys())
    if missing:
        raise ValueError(f"Feedback-Datensatz fehlt: {', '.join(missing)}")

    if not isinstance(payload["bewertung"], dict) or not payload["bewertung"]:
        raise ValueError("Feld 'bewertung' muss ein nicht-leeres Objekt sein.")

    note_payload = payload["notenempfehlung"]
    note_required = {"durchschnitt", "note", "bezeichnung", "begruendung"}
    if not isinstance(note_payload, dict):
        raise ValueError("Feld 'notenempfehlung' muss ein Objekt sein.")
    note_missing = sorted(note_required - note_payload.keys())
    if note_missing:
        raise ValueError(f"Notenempfehlung fehlt: {', '.join(note_missing)}")

    if not isinstance(note_payload["durchschnitt"], (int, float)):
        raise ValueError("'notenempfehlung.durchschnitt' muss numerisch sein.")
    if not isinstance(note_payload["note"], int):
        raise ValueError("'notenempfehlung.note' muss eine ganze Zahl sein.")
    if not isinstance(note_payload["bezeichnung"], str) or not isinstance(
        note_payload["begruendung"], str
    ):
        raise ValueError(
            "'notenempfehlung.bezeichnung' und 'begruendung' muessen Strings sein."
        )

    criteria = [
        parse_criterion(key, value) for key, value in payload["bewertung"].items()
    ]

    return FeedbackData(
        datei=str(payload["datei"]),
        schueler=str(payload["schueler"]) if payload.get("schueler") else None,
        klasse=str(payload["klasse"]) if payload.get("klasse") else None,
        textsorte=str(payload["textsorte"]),
        fach=str(payload["fach"]),
        schulstufe=str(payload["schulstufe"]),
        rubrik=str(payload["rubrik"]),
        bewertung=criteria,
        notenempfehlung=GradeRecommendation(
            durchschnitt=float(note_payload["durchschnitt"]),
            note=note_payload["note"],
            bezeichnung=note_payload["bezeichnung"],
            begruendung=note_payload["begruendung"],
        ),
        hinweise=ensure_list(payload.get("hinweise"), "hinweise"),
    )


def load_feedback_json(path: Path) -> FeedbackData:
    """Laedt und validiert eine JSON-Datei."""

    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("Die JSON-Wurzel muss ein Objekt sein.")
    return parse_feedback_data(data)


def criterion_label(data: FeedbackData, key: str) -> str:
    """Liefert den Anzeige-Namen eines Kriteriums."""

    german_labels = {
        "inhalt": "INHALT",
        "textstruktur": "TEXTSTRUKTUR",
        "ausdruck": "STIL UND AUSDRUCK"
        if data.schulstufe == "Oberstufe"
        else "AUSDRUCK",
        "sprachrichtigkeit": "NORMATIVE SPRACHRICHTIGKEIT"
        if data.schulstufe == "Oberstufe"
        else "SPRACHRICHTIGKEIT",
    }
    english_labels = {
        "task_achievement": "TASK ACHIEVEMENT",
        "organisation_layout": "ORGANISATION AND LAYOUT",
        "lexical_range_accuracy": "LEXICAL RANGE AND ACCURACY",
        "grammatical_range_accuracy": "GRAMMATICAL RANGE AND ACCURACY",
    }
    labels = german_labels if data.fach == "Deutsch" else english_labels
    return labels.get(key, key.replace("_", " ").upper())


def ordered_criteria(data: FeedbackData) -> list[CriterionFeedback]:
    """Sortiert Kriterien in einer fachlich passenden Reihenfolge."""

    desired = GERMAN_ORDER if data.fach == "Deutsch" else ENGLISH_ORDER
    rank = {name: index for index, name in enumerate(desired)}
    return sorted(data.bewertung, key=lambda item: rank.get(item.key, len(rank)))


def render_list_section(
    doc: Document,
    title: str,
    entries: list[str] | None,
    color: RGBColor | None = None,
) -> None:
    """Rendert eine farbige Ueberschrift plus Liste oder Platzhalter."""

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    if color is not None:
        run.font.color.rgb = color
    paragraph.paragraph_format.space_after = Pt(1)
    if entries:
        for entry in entries:
            add_bullet(doc, entry, color=color)
    else:
        add_bullet(doc, "Keine Angaben.")


def add_summary_table(doc: Document, data: FeedbackData) -> None:
    """Fuegt eine Uebersichtstabelle mit allen Kriterien und der Gesamtnote ein."""

    criteria = ordered_criteria(data)
    note = data.notenempfehlung.note
    grade_color = _note_color(note)

    table = doc.add_table(rows=1 + len(criteria) + 1, cols=3)
    table.style = "Table Grid"

    for cell, text in zip(
        table.rows[0].cells, ["Kriterium", "Stufe / Bewertung", "Punkte"]
    ):
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _set_cell_bg(cell, "1F497D")

    for row_idx, crit in enumerate(criteria, start=1):
        cells = table.rows[row_idx].cells
        label = criterion_label(data, crit.key)
        for cell, text in zip(cells, [label, crit.stufe, f"{crit.punkte:g}"]):
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        if row_idx % 2 == 0:
            for cell in cells:
                _set_cell_bg(cell, "EEF2FF")

    grade_cells = table.rows[-1].cells
    for cell, text in zip(
        grade_cells,
        ["GESAMTNOTE", "", f"{note} – {data.notenempfehlung.bezeichnung}"],
    ):
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        _set_cell_bg(cell, str(grade_color))

    doc.add_paragraph()


def build_feedback_document(
    data: FeedbackData, config: dict[str, Any] | None = None
) -> Document:
    """Erstellt das DOCX-Dokument aus dem validierten Datensatz."""

    doc = Document()
    setup_page(doc, config)
    add_document_header(doc, data, config)

    add_summary_table(doc, data)
    add_divider(doc)

    if data.hinweise:
        add_section_header(doc, "HINWEISE")
        for hinweis in data.hinweise:
            add_bullet(doc, hinweis)
        add_divider(doc)

    for index, criterion in enumerate(ordered_criteria(data), start=1):
        add_section_header(doc, f"{index}. {criterion_label(data, criterion.key)}")
        points_text = f"{criterion.punkte:g} Punkte"
        if criterion.gewicht is not None:
            points_text = f"{points_text} | Gewicht: {criterion.gewicht:g} %"
        add_label(doc, "Bewertung", f"{criterion.stufe} [{points_text}]")
        doc.add_paragraph()

        render_list_section(doc, "Staerken:", criterion.staerken, color=C_STRENGTH)
        doc.add_paragraph()

        if criterion.key == "sprachrichtigkeit" and criterion.fehler_detail:
            render_list_section(
                doc, "Fehler im Detail:", criterion.fehler_detail, color=C_WEAKNESS
            )
            doc.add_paragraph()
            render_list_section(
                doc,
                "Fehlerschwerpunkte:",
                criterion.fehlerschwerpunkte or [],
                color=C_WEAKNESS,
            )
            doc.add_paragraph()
        elif criterion.key == "grammatical_range_accuracy" and criterion.fehler_detail:
            render_list_section(
                doc, "Fehler im Detail:", criterion.fehler_detail, color=C_WEAKNESS
            )
            doc.add_paragraph()

        if criterion.rhetorische_figuren:
            render_list_section(
                doc, "Rhetorische Figuren:", criterion.rhetorische_figuren
            )
            doc.add_paragraph()

        render_list_section(
            doc, "Schwaechen / Fehler:", criterion.schwaechen, color=C_WEAKNESS
        )
        doc.add_paragraph()
        render_list_section(
            doc, "Verbesserungsvorschlaege:", criterion.vorschlaege, color=C_SUGGESTION
        )
        add_divider(doc)

    add_section_header(doc, "NOTENEMPFEHLUNG")
    add_divider(doc)
    for criterion in ordered_criteria(data):
        label = criterion_label(data, criterion.key).title()
        value = f"{criterion.stufe} ({criterion.punkte:g} Punkte)"
        if criterion.gewicht is not None:
            value = f"{value} x {criterion.gewicht:g} %"
        add_label(doc, label, value)

    doc.add_paragraph()
    add_label(doc, "Durchschnitt", f"{data.notenempfehlung.durchschnitt:.2f}")

    note_paragraph = doc.add_paragraph()
    note_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_run = note_paragraph.add_run(
        f"Empfohlene Note: {data.notenempfehlung.note} - {data.notenempfehlung.bezeichnung}"
    )
    note_run.bold = True
    note_run.font.size = Pt(12)
    note_run.font.color.rgb = _note_color(data.notenempfehlung.note)

    doc.add_paragraph()
    add_body(doc, "Begruendung:")
    add_body(doc, data.notenempfehlung.begruendung)
    doc.add_paragraph()
    add_body(
        doc,
        "HINWEIS: Diese Notenempfehlung ist ein Hilfsmittel und ersetzt nicht die paedagogische Beurteilung durch die Lehrerin.",
    )
    add_divider(doc)
    return doc


def build_statistics_document(
    stats: dict[str, Any],
    config: dict[str, Any] | None = None,
    klasse_name: str = "",
) -> Document:
    """Erstellt ein A4-DOCX-Dokument mit Klassen-Statistiken.

    Args:
        stats: Rückgabewert von compute_statistics()
        config: saga_config.toml-Dict (optional, für Header/Footer)
        klasse_name: Anzeigename der Klasse (optional)
    """
    _NOTE_LABELS_MAP = {
        1: "Sehr gut",
        2: "Gut",
        3: "Befriedigend",
        4: "Genügend",
        5: "Nicht gen.",
    }
    doc = Document()
    setup_page(doc, config)
    add_heading(doc, "KLASSENAUSWERTUNG", level=1)

    # Meta-Tabelle
    meta = doc.add_table(rows=0, cols=2)
    meta.autofit = True
    today = date.today().strftime("%d.%m.%Y")
    meta_rows: list[tuple[str, str]] = []
    if klasse_name:
        meta_rows.append(("Klasse", klasse_name))
    meta_rows.append(("Datum", today))
    docx_cfg = (config or {}).get("docx", {})
    teacher = docx_cfg.get("teacher_name", "")
    school = docx_cfg.get("school_name", "")
    if teacher:
        meta_rows.append(("Lehrer/in", teacher))
    if school:
        meta_rows.append(("Schule", school))
    for label, value in meta_rows:
        row = meta.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

    # Übersicht
    total = stats.get("total", 0)
    avg = stats.get("grade_average", 0.0)
    p_total = doc.add_paragraph()
    r_total = p_total.add_run(f"Ausgewertete Schüler/innen: {total}")
    r_total.bold = True
    r_total.font.size = Pt(11)
    if total > 0:
        p_avg = doc.add_paragraph()
        r_avg = p_avg.add_run(f"Gesamtdurchschnitt: {avg:.2f}")
        r_avg.bold = True
        r_avg.font.size = Pt(11)
        r_avg.font.color.rgb = _note_color(max(1, min(5, round(avg))))
    doc.add_paragraph()

    # Notenverteilung
    add_heading(doc, "Notenverteilung", level=2)
    dist = stats.get("grade_distribution", {})
    grade_table = doc.add_table(rows=1 + 5, cols=3)
    grade_table.autofit = True
    hrow = grade_table.rows[0]
    hrow.cells[0].text = "Note"
    hrow.cells[1].text = "Bezeichnung"
    hrow.cells[2].text = f"Anzahl (von {total})"
    for cell in hrow.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = C_PRIMARY
    for i, note in enumerate(range(1, 6), 1):
        count = dist.get(note, 0)
        pct = (count / total * 100) if total > 0 else 0.0
        row = grade_table.rows[i]
        row.cells[0].text = str(note)
        row.cells[1].text = _NOTE_LABELS_MAP.get(note, "")
        row.cells[2].text = f"{count}  ({pct:.1f}%)"
        color = _note_color(note)
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = color
    doc.add_paragraph()

    # Kriterien-Auswertung
    crit_avgs = stats.get("criteria_averages", {})
    if crit_avgs:
        add_heading(doc, "Kriterien-Auswertung", level=2)
        weakest = stats.get("weakest_criterion")
        strongest = stats.get("strongest_criterion")
        crit_table = doc.add_table(rows=1 + len(crit_avgs), cols=4)
        crit_table.autofit = True
        hrow = crit_table.rows[0]
        hrow.cells[0].text = "Kriterium"
        hrow.cells[1].text = "Ø Punkte"
        hrow.cells[2].text = "Min"
        hrow.cells[3].text = "Max"
        for cell in hrow.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = C_PRIMARY
        for i, (key, vals) in enumerate(
            sorted(crit_avgs.items(), key=lambda x: x[1]["avg"]), 1
        ):
            row = crit_table.rows[i]
            label = key.replace("_", " ").title()
            row.cells[0].text = label
            row.cells[1].text = f"{vals['avg']:.2f}"
            row.cells[2].text = f"{vals['min']:.1f}"
            row.cells[3].text = f"{vals['max']:.1f}"
            if key == weakest:
                highlight = C_GRADE_5
            elif key == strongest:
                highlight = C_GRADE_1
            else:
                highlight = None
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        if highlight:
                            run.font.color.rgb = highlight
                            run.bold = True

    return doc


def output_filename(original_name: str) -> str:
    """Leitet den Zieldateinamen aus der Originaldatei ab."""

    return f"{Path(original_name).stem}_feedback.docx"


def log_error(log_path: Path, message: str) -> None:
    """Schreibt eine Fehlermeldung ins Fehlerlog."""

    log_path.parent.mkdir(parents=True, exist_ok=True)
    with log_path.open("a", encoding="utf-8") as handle:
        handle.write(f"{message}\n")


def collect_json_files(paths: ProjectPaths, selected_file: str | None) -> list[Path]:
    """Bestimmt die zu verarbeitenden JSON-Dateien."""

    if selected_file:
        candidate = Path(selected_file)
        if not candidate.is_absolute():
            candidate = paths.feedback_data_dir / selected_file
        if not candidate.exists():
            raise FileNotFoundError(f"JSON-Datei nicht gefunden: {candidate}")
        return [candidate]
    return sorted(paths.feedback_data_dir.glob("*.json"))


def process_file(
    json_path: Path, paths: ProjectPaths, force: bool = False, dry_run: bool = False
) -> str:
    """Verarbeitet eine einzelne JSON-Datei."""

    feedback = load_feedback_json(json_path)
    output_path = paths.output_dir / output_filename(feedback.datei)

    if output_path.exists() and not force:
        return f"Uebersprungen (existiert bereits): {output_path.name}"

    if dry_run:
        return f"Dry-run: wuerde {json_path.name} -> {output_path.name} verarbeiten"

    paths.output_dir.mkdir(parents=True, exist_ok=True)
    document = build_feedback_document(feedback)
    document.save(output_path)
    return f"Gespeichert: {output_path.name}"


def main() -> int:
    """CLI-Einstiegspunkt."""

    args = parse_args()
    paths = project_paths()
    paths.feedback_data_dir.mkdir(parents=True, exist_ok=True)
    paths.output_dir.mkdir(parents=True, exist_ok=True)

    if not args.dry_run:
        paths.fehlerlog.write_text("", encoding="utf-8")

    try:
        json_files = collect_json_files(paths, args.file)
    except Exception as exc:
        log_error(paths.fehlerlog, str(exc))
        print(str(exc))
        return 1

    if not json_files:
        print("Keine JSON-Dateien in /output/feedback_data gefunden.")
        return 0

    failures = 0
    for json_path in json_files:
        try:
            message = process_file(
                json_path, paths, force=args.force, dry_run=args.dry_run
            )
            print(message)
        except Exception as exc:
            failures += 1
            message = f"{json_path.name} - Fehler: {exc}"
            log_error(paths.fehlerlog, message)
            print(message)

    if failures:
        print(f"Abgeschlossen mit {failures} Fehler(n).")
        return 1

    print("Fertig.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
