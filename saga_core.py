#!/usr/bin/env python3
"""
SAGA Core – Gemeinsame Logik-Funktionen fuer Wizard und Dashboard.
"""

from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
import threading
import time
import tomllib
import tomlkit
import urllib.error
import urllib.request
from pathlib import Path
from typing import Any

if sys.version_info < (3, 11):
    print("Python 3.11+ wird benoetigt.")
    raise SystemExit(1)

try:
    from docx import Document as DocxDocument
except ImportError:
    print("python-docx fehlt: pip install python-docx")
    raise SystemExit(1)

try:
    import jsonschema as _jsmod
except ImportError:
    _jsmod = None

sys.path.insert(0, str(Path(__file__).resolve().parent))
import generate_feedback as gf

PROJECT_ROOT = Path(__file__).resolve().parent

VERSION = "0.7.0"


def _load_dotenv() -> None:
    """Lädt .env aus dem Projektverzeichnis — überschreibt keine bereits gesetzten Variablen."""
    env_path = PROJECT_ROOT / ".env"
    if not env_path.exists():
        return
    try:
        from dotenv import load_dotenv

        load_dotenv(env_path, override=False)
        return
    except ImportError:
        pass
    # Fallback: manuelles Parsen ohne externe Abhängigkeit
    with env_path.open(encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, _, value = line.partition("=")
            key = key.strip()
            value = value.strip().strip('"').strip("'")
            if key and key not in os.environ:
                os.environ[key] = value


def load_config() -> dict[str, Any]:
    _load_dotenv()
    config_path = PROJECT_ROOT / "saga_config.toml"
    if not config_path.exists():
        raise FileNotFoundError(f"Konfigurationsdatei nicht gefunden: {config_path}")
    with config_path.open("rb") as f:
        return tomllib.load(f)


def _load_toml_doc() -> tomlkit.TOMLDocument:
    """Lädt saga_config.toml als bearbeitbares tomlkit-Dokument (Kommentare bleiben erhalten)."""
    return tomlkit.parse((PROJECT_ROOT / "saga_config.toml").read_text(encoding="utf-8"))


def _save_toml_doc(doc: tomlkit.TOMLDocument) -> None:
    """Schreibt ein tomlkit-Dokument zurück nach saga_config.toml."""
    (PROJECT_ROOT / "saga_config.toml").write_text(tomlkit.dumps(doc), encoding="utf-8")


def resolve_path(config: dict[str, Any], key: str) -> Path:
    return PROJECT_ROOT / config["paths"][key]


def count_words(docx_path: Path) -> int:
    try:
        doc = DocxDocument(str(docx_path))
        return sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
    except Exception:
        return 0


def read_docx_text(docx_path: Path) -> str:
    doc = DocxDocument(str(docx_path))
    return "\n".join(p.text for p in doc.paragraphs)


def read_docx_rich(docx_path: Path) -> str:
    """Return the full document text with basic Rich markup for terminal display.

    Paragraph separation is the top priority; headings get bold+underline,
    inline bold/italic are also rendered.  Empty paragraphs are collapsed so
    that a single blank line always separates paragraphs.
    """
    doc = DocxDocument(str(docx_path))
    lines: list[str] = []
    prev_empty = False

    for para in doc.paragraphs:
        raw = para.text

        # Completely empty paragraph → use as separator (max one blank line)
        if not raw.strip():
            if not prev_empty and lines:
                lines.append("")
            prev_empty = True
            continue
        prev_empty = False

        # Detect heading style
        style_name = (para.style.name or "").lower()
        is_heading = style_name.startswith("heading") or style_name.startswith("überschrift")

        # Build the line with per-run inline markup
        parts: list[str] = []
        for run in para.runs:
            if not run.text:
                continue
            # Escape Rich markup characters in the raw text
            text = run.text.replace("[", r"\[")
            if run.bold and run.italic:
                text = f"[bold italic]{text}[/bold italic]"
            elif run.bold:
                text = f"[bold]{text}[/bold]"
            elif run.italic:
                text = f"[italic]{text}[/italic]"
            parts.append(text)

        line = "".join(parts) if parts else raw.replace("[", r"\[")

        if is_heading:
            line = f"[bold underline]{line}[/bold underline]"

        lines.append(line)

    return "\n".join(lines)


def load_rubric(rubric_filename: str, config: dict[str, Any]) -> str:
    rubric_dir = resolve_path(config, "rubrics")
    rubric_path = rubric_dir / rubric_filename
    if not rubric_path.exists():
        raise FileNotFoundError(f"Rubrik nicht gefunden: {rubric_path}")
    return rubric_path.read_text(encoding="utf-8")


def list_all_rubrics(config: dict[str, Any]) -> list[str]:
    """Gibt alle .md-Dateien aus rubrics/ zurück (ohne README-Dateien)."""
    rubric_dir = resolve_path(config, "rubrics")
    if not rubric_dir.exists():
        return []
    return sorted(
        f.name for f in rubric_dir.glob("*.md") if not f.name.upper().startswith("README")
    )


def load_rubric_for_aufgabe(config: dict[str, Any], klasse: str | None, aufgabe: str | None) -> str:
    """Lädt die Rubrik für eine Aufgabe — mit Fallback-Kette.

    1. aufgabe_cfg["rubric"] (Dateiname in rubrics/)
    2. default_rubric_for(fach, schulstufe)
    3. Erste verfügbare .md-Datei in rubrics/
    """
    auf_cfg = get_aufgabe_cfg(config, klasse or "", aufgabe or "") if klasse and aufgabe else {}
    rubric_name = auf_cfg.get("rubric", "")
    if rubric_name:
        try:
            return load_rubric(rubric_name, config)
        except FileNotFoundError:
            pass
    fach = auf_cfg.get("fach") or config.get("defaults", {}).get("fach", "Deutsch")
    schulstufe = auf_cfg.get("schulstufe") or config.get("defaults", {}).get(
        "schulstufe", "Oberstufe"
    )
    default = default_rubric_for(fach, schulstufe, config)
    if default:
        return load_rubric(default, config)
    rubric_dir = resolve_path(config, "rubrics")
    for f in sorted(rubric_dir.glob("*.md")):
        if not f.name.upper().startswith("README"):
            return f.read_text(encoding="utf-8")
    raise FileNotFoundError("Keine Rubrik gefunden")


def set_rubric_for_aufgabe(klasse: str, aufgabe: str, rubric_name: str) -> None:
    """Setzt die Rubrik einer Aufgabe in saga_config.toml (kein Dateikopieren)."""
    doc = _load_toml_doc()
    doc["classes"][klasse]["aufgaben"][aufgabe]["rubric"] = rubric_name
    _save_toml_doc(doc)


def attach_rubric_to_aufgabe(klasse: str, aufgabe: str, source_path: Path) -> str:
    """Kopiert source_path nach rubrics/{aufgabe}_{dateiname} und aktualisiert die Config.

    Gibt den neuen Dateinamen (relativ zu rubrics/) zurück.
    """
    config = load_config()
    rubric_dir = resolve_path(config, "rubrics")
    rubric_dir.mkdir(parents=True, exist_ok=True)
    dest_name = f"{aufgabe}_{source_path.name}"
    dest = rubric_dir / dest_name
    shutil.copy2(str(source_path), str(dest))
    set_rubric_for_aufgabe(klasse, aufgabe, dest_name)
    return dest_name


def load_schema(config: dict[str, Any]) -> dict[str, Any]:
    schema_path = resolve_path(config, "schema")
    if not schema_path.exists():
        return {}
    return json.loads(schema_path.read_text(encoding="utf-8"))


def load_example_fixture() -> str:
    fixtures_dir = PROJECT_ROOT / "tests" / "fixtures"
    candidates = sorted(fixtures_dir.glob("*.json"))
    if not candidates:
        return "{}"
    return candidates[0].read_text(encoding="utf-8")


def build_analysis_prompt(
    docx_text: str,
    rubric_content: str,
    fach: str,
    schulstufe: str,
    textsorte: str,
    config: dict[str, Any],
    schueler: str = "",
) -> str:
    schema = load_schema(config)
    example = load_example_fixture()
    schema_str = json.dumps(schema, indent=2, ensure_ascii=False)
    example_str = (
        json.dumps(json.loads(example), indent=2, ensure_ascii=False) if example != "{}" else ""
    )

    return (
        "Du bist ein Korrekturassistent für österreichische Gymnasium-Schularbeiten.\n"
        f"Fach: {fach}\nSchulstufe: {schulstufe}\nTextsorte: {textsorte}\n"
        + (f"Schüler/in: {schueler}\n" if schueler else "")
        + "\n"
        "BEWERTUNGSRASTER:\n---\n"
        f"{rubric_content}\n---\n\n"
        "SCHÜLERTEXT:\n---\n"
        f"{docx_text}\n---\n\n"
        "JSON-SCHEMA (dein Output MUSS konform sein):\n---\n"
        f"{schema_str}\n---\n\n"
        "BEISPIEL-JSON:\n---\n"
        f"{example_str}\n---\n\n"
        "AUFGABE:\n"
        "Analysiere den Schülertext anhand des Bewertungsrasters.\n"
        "Erstelle eine Bewertung für jedes Kriterium mit Stufe, Punkten, Stärken, Schwächen und Vorschlägen.\n"
        "Berechne eine Notenempfehlung.\n\n"
        "WICHTIG: Antworte NUR mit validem JSON. Kein Markdown, kein Erklärtext, keine ```json-Blöcke.\n"
        "Das JSON muss dem obigen Schema entsprechen.\n"
    )


def extract_json_from_llm(text: str) -> dict[str, Any]:
    fenced = re.search(r"```(?:json)?\s*\n?(.*?)```", text, re.DOTALL)
    candidate = fenced.group(1).strip() if fenced else text.strip()
    brace_match = re.search(r"\{.*\}", candidate, re.DOTALL)
    if brace_match:
        candidate = brace_match.group(0)
    return json.loads(candidate)


def validate_against_schema(data: dict[str, Any], schema: dict[str, Any]) -> list[str]:
    if _jsmod is None or not schema:
        return []
    errors: list[str] = []
    for err in sorted(
        _jsmod.Draft202012Validator(schema).iter_errors(data), key=lambda e: e.message
    ):
        errors.append(err.message)
    return errors


# ---------------------------------------------------------------------------
# Robuste LLM-Analyse-Pipeline: Prompt → API → JSON → Validierung → Retry
# ---------------------------------------------------------------------------

_MAX_RETRIES = 3
_RETRY_BACKOFF_SECONDS = 2


def _build_retry_prompt(
    original_prompt: str,
    error_message: str,
    llm_raw_response: str,
    attempt: int,
) -> str:
    """Erzeugt einen Retry-Prompt, der das LLM zur JSON-Korrektur auffordert."""
    return (
        f"DEINE VORHERIGE ANTWORT WAR UNGUELTIG (Versuch {attempt}).\n\n"
        f"FEHLER: {error_message}\n\n"
        f"DEINE ANTWORT:\n---\n{llm_raw_response[:2000]}\n---\n\n"
        f"ORIGINAL-AUFGABE (zur Erinnerung):\n---\n{original_prompt[:3000]}\n---\n\n"
        f"ANTWORTE JETZT AUSSCHLIESSLICH MIT GUELTIGEM JSON. "
        f"Kein Markdown, kein Code-Block, kein Erklaertext. "
        f"NUR das JSON-Objekt, beginnend mit {{ und endend mit }}."
    )


def run_llm_analysis(
    docx_text: str,
    rubric_content: str,
    fach: str,
    schulstufe: str,
    textsorte: str,
    config: dict[str, Any],
    schueler: str = "",
    cancel_event: threading.Event | None = None,
    max_retries: int = _MAX_RETRIES,
) -> tuple[dict[str, Any] | None, list[str]]:
    """
    Fuehrt die vollstaendige LLM-Analyse durch: Prompt bauen, API aufrufen,
    JSON extrahieren, gegen Schema validieren. Bei Fehlern wird bis zu
    *max_retries*-mal wiederholt.

    Returns:
        (data, errors) – data ist das validierte JSON-Dict oder None,
        errors ist eine Liste aller aufgetretenen Fehlermeldungen.
    """
    schema = load_schema(config)
    original_prompt = build_analysis_prompt(
        docx_text, rubric_content, fach, schulstufe, textsorte, config, schueler
    )
    errors: list[str] = []
    current_prompt = original_prompt

    for attempt in range(1, max_retries + 1):
        if cancel_event is not None and cancel_event.is_set():
            errors.append("Analyse abgebrochen")
            return None, errors

        raw_response = run_llm_api(current_prompt, config, cancel_event=cancel_event, schema=schema)

        if cancel_event is not None and cancel_event.is_set():
            errors.append("Analyse abgebrochen")
            return None, errors

        if raw_response.startswith("FEHLER"):
            errors.append(f"API-Fehler (Versuch {attempt}): {raw_response}")
            # Bei API-Fehlern lohnt sich kein Retry mit gleichem Prompt
            return None, errors

        # Versuch, JSON zu extrahieren
        try:
            data = extract_json_from_llm(raw_response)
        except (json.JSONDecodeError, AttributeError) as e:
            errors.append(f"JSON-Extraktion fehlgeschlagen (Versuch {attempt}): {e}")
            current_prompt = _build_retry_prompt(
                original_prompt, f"Kein gueltiges JSON gefunden: {e}", raw_response, attempt
            )
            time.sleep(_RETRY_BACKOFF_SECONDS * attempt)
            continue

        # Gegen Schema validieren
        validation_errors = validate_against_schema(data, schema)
        if validation_errors:
            errors.append(
                f"Schema-Validierung fehlgeschlagen (Versuch {attempt}): "
                f"{'; '.join(validation_errors[:3])}"
            )
            current_prompt = _build_retry_prompt(
                original_prompt,
                f"Schema-Verletzung: {'; '.join(validation_errors[:3])}",
                raw_response,
                attempt,
            )
            time.sleep(_RETRY_BACKOFF_SECONDS * attempt)
            continue

        # Erfolg
        return data, []

    errors.append(
        f"Analyse nach {max_retries} Versuchen fehlgeschlagen. Siehe fehlerlog.txt fuer Details."
    )
    return None, errors


def build_project_paths(
    config: dict[str, Any],
    klasse: str | None = None,
    aufgabe: str | None = None,
) -> gf.ProjectPaths:
    """Baut Projektpfade für die angegebene Klasse + Aufgabe.

    Hierarchie: Aufgabe-Pfade > Klassen-Pfade > [paths]-Fallback.
    Wenn klasse/aufgabe None sind, werden aktive Werte aus config gelesen.
    """
    classes_cfg = config.get("classes", {})
    class_names = [k for k in classes_cfg if k != "active" and isinstance(classes_cfg[k], dict)]

    if class_names:
        if klasse is None:
            klasse = classes_cfg.get("active", class_names[0])
        cls = classes_cfg.get(klasse) if klasse else None
        if isinstance(cls, dict):
            # Aufgaben-Pfade ermitteln
            if aufgabe is None:
                aufgabe = cls.get("active_aufgabe")
            auf_cfg: dict[str, Any] = {}
            if aufgabe:
                auf_cfg = cls.get("aufgaben", {}).get(aufgabe, {})

            input_rel = auf_cfg.get("input") or cls.get("input", "input")
            output_rel = auf_cfg.get("output") or cls.get("output", "output")
            output_dir = PROJECT_ROOT / output_rel
            return gf.ProjectPaths(
                root=PROJECT_ROOT,
                input_dir=PROJECT_ROOT / input_rel,
                output_dir=output_dir,
                feedback_data_dir=output_dir / "feedback_data",
                fehlerlog=output_dir / "fehlerlog.txt",
            )

    # Fallback: klassischer [paths]-Block
    output_dir = resolve_path(config, "output")
    return gf.ProjectPaths(
        root=PROJECT_ROOT,
        input_dir=resolve_path(config, "input"),
        output_dir=output_dir,
        feedback_data_dir=resolve_path(config, "feedback_data"),
        fehlerlog=resolve_path(config, "fehlerlog"),
    )


def list_classes(config: dict[str, Any]) -> list[str]:
    """Gibt alle konfigurierten Klassennamen zurück (ohne 'active')."""
    classes_cfg = config.get("classes", {})
    return [k for k in classes_cfg if k != "active" and isinstance(classes_cfg[k], dict)]


def active_klasse(config: dict[str, Any]) -> str | None:
    """Gibt den aktiven Klassennamen zurück, oder None wenn keine Klassen konfiguriert."""
    names = list_classes(config)
    if not names:
        return None
    return config.get("classes", {}).get("active", names[0])


def save_active_klasse(klasse: str) -> None:
    """Schreibt die aktive Klasse in saga_config.toml."""
    doc = _load_toml_doc()
    doc["classes"]["active"] = klasse
    _save_toml_doc(doc)


def add_class_to_config(name: str, input_rel: str, output_rel: str) -> None:
    """Fügt eine neue Klasse in saga_config.toml ein und setzt sie als aktiv."""
    doc = _load_toml_doc()
    doc["classes"]["active"] = name
    new_cls = tomlkit.table()
    new_cls.add("input", input_rel)
    new_cls.add("output", output_rel)
    doc["classes"].add(name, new_cls)
    _save_toml_doc(doc)


def list_aufgaben(config: dict[str, Any], klasse: str) -> list[str]:
    """Gibt Slug-Liste aller Aufgaben einer Klasse zurück (Reihenfolge wie in config)."""
    cls = config.get("classes", {}).get(klasse, {})
    aufgaben = cls.get("aufgaben", {})
    return [k for k in aufgaben if isinstance(aufgaben[k], dict)]


def active_aufgabe(config: dict[str, Any], klasse: str) -> str | None:
    """Gibt die aktive Aufgabe einer Klasse zurück, oder None wenn keine vorhanden."""
    names = list_aufgaben(config, klasse)
    if not names:
        return None
    cls = config.get("classes", {}).get(klasse, {})
    return cls.get("active_aufgabe", names[-1])


def get_aufgabe_cfg(config: dict[str, Any], klasse: str, aufgabe: str) -> dict[str, Any]:
    """Gibt die Konfigurations-Dict einer Aufgabe zurück (oder {})."""
    cls = config.get("classes", {}).get(klasse, {})
    return cls.get("aufgaben", {}).get(aufgabe, {})


def aufgabe_defaults(
    config: dict[str, Any], klasse: str | None, aufgabe: str | None
) -> dict[str, str]:
    """Gibt fach/schulstufe/textsorte/rubric der Aufgabe zurück.

    Leere Strings wenn nicht konfiguriert — Caller fällt dann auf globale Defaults zurück.
    """
    if not klasse or not aufgabe:
        return {}
    auf_cfg = get_aufgabe_cfg(config, klasse, aufgabe)
    return {
        "fach": auf_cfg.get("fach", ""),
        "schulstufe": auf_cfg.get("schulstufe", ""),
        "textsorte": auf_cfg.get("textsorte", ""),
        "rubric": auf_cfg.get("rubric", ""),
    }


def save_active_aufgabe(klasse: str, aufgabe: str) -> None:
    """Schreibt active_aufgabe für die Klasse in saga_config.toml."""
    doc = _load_toml_doc()
    doc["classes"][klasse]["active_aufgabe"] = aufgabe
    _save_toml_doc(doc)


def add_aufgabe_to_config(
    klasse: str,
    slug: str,
    label: str,
    fach: str,
    schulstufe: str,
    textsorte: str,
    rubric: str,
) -> None:
    """Fügt eine neue Aufgabe zur Klasse in saga_config.toml hinzu.

    Setzt die neue Aufgabe sofort als active_aufgabe und erstellt die Ordner.
    """
    doc = _load_toml_doc()
    cls_doc = doc["classes"][klasse]
    cls_doc["active_aufgabe"] = slug

    base_input = str(cls_doc.get("input", f"input/{klasse}"))
    base_output = str(cls_doc.get("output", f"output/{klasse}"))
    input_rel = f"{base_input}/{slug}"
    output_rel = f"{base_output}/{slug}"

    if "aufgaben" not in cls_doc:
        cls_doc.add("aufgaben", tomlkit.table())

    auf_tbl = tomlkit.table()
    auf_tbl.add("label", label)
    auf_tbl.add("fach", fach)
    auf_tbl.add("schulstufe", schulstufe)
    auf_tbl.add("textsorte", textsorte)
    auf_tbl.add("rubric", rubric)
    auf_tbl.add("input", input_rel)
    auf_tbl.add("output", output_rel)
    cls_doc["aufgaben"].add(slug, auf_tbl)

    _save_toml_doc(doc)

    # Ordner anlegen
    (PROJECT_ROOT / input_rel).mkdir(parents=True, exist_ok=True)
    (PROJECT_ROOT / output_rel).mkdir(parents=True, exist_ok=True)


def log_tui_error(paths: gf.ProjectPaths, message: str) -> None:
    paths.fehlerlog.parent.mkdir(parents=True, exist_ok=True)
    with paths.fehlerlog.open("a", encoding="utf-8") as f:
        f.write(f"[TUI] {message}\n")


def rubric_options_for(fach: str, schulstufe: str, config: dict[str, Any]) -> list[str]:
    rubric_dir = resolve_path(config, "rubrics")
    if fach == "Deutsch":
        if schulstufe == "Oberstufe":
            files = ["srdp_deutsch_oberstufe.md"]
        else:
            files = ["deutsch_unterstufe.md"]
    else:
        if schulstufe == "Unterstufe":
            files = ["englisch_a2.md"]
        else:
            files = ["srdp_englisch_b2.md", "srdp_englisch_b1.md"]
    return [f for f in files if (rubric_dir / f).exists()]


def default_rubric_for(fach: str, schulstufe: str, config: dict[str, Any]) -> str:
    mapping = config.get("rubric_mapping", {})
    configured = mapping.get(f"{fach}+{schulstufe}")
    options = rubric_options_for(fach, schulstufe, config)
    if configured in options:
        return configured
    return options[0] if options else ""


def copy_to_clipboard(text: str) -> bool:
    for cmd in [
        "clip.exe",
        "xclip -selection clipboard",
        "xsel --clipboard",
        "wl-copy",
    ]:
        parts = cmd.split()
        try:
            proc = subprocess.run(
                parts,
                input=text.encode(),
                timeout=5,
                capture_output=True,
            )
            if proc.returncode == 0:
                return True
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    return False


def save_settings(fach: str, schulstufe: str, provider: str, model: str) -> None:
    """Schreibt API- und Default-Einstellungen in saga_config.toml."""
    doc = _load_toml_doc()
    doc["defaults"]["fach"] = fach
    doc["defaults"]["schulstufe"] = schulstufe
    doc["api"]["provider"] = provider
    doc["api"]["model"] = model
    _save_toml_doc(doc)


def open_file(path: Path) -> bool:
    """Öffnet eine Datei mit der systemseitigen Standard-App.

    Probiert wslview (WSL), xdg-open (Linux) und open (macOS) der Reihe nach.
    Gibt True zurück wenn ein Kommando gestartet wurde, False wenn keines gefunden.
    """
    path = path.resolve()
    candidates = [
        ["wslview", str(path)],
        ["xdg-open", str(path)],
        ["open", str(path)],
    ]
    for cmd in candidates:
        try:
            subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            return True
        except FileNotFoundError:
            continue
    return False


def run_agent_sync(cmd_template: str, prompt: str, timeout: int) -> str:
    parts = cmd_template.split()
    try:
        result = subprocess.run(
            parts,
            input=prompt,
            capture_output=True,
            text=True,
            timeout=timeout,
        )
        if result.returncode != 0:
            return f"FEHLER (Exit {result.returncode}): {result.stderr[:500]}"
        return result.stdout
    except subprocess.TimeoutExpired:
        return f"FEHLER: Timeout nach {timeout}s"
    except FileNotFoundError as e:
        return f"FEHLER: Befehl nicht gefunden: {e}"


def _call_ollama_native(
    base_url: str,
    model: str,
    prompt: str,
    timeout: int,
    cancel_event: threading.Event | None = None,
) -> str:
    """Ruft Ollama über die native /api/chat-Schnittstelle auf."""
    if cancel_event is not None and cancel_event.is_set():
        return "FEHLER: Abgebrochen"
    url = base_url.rstrip("/") + "/api/chat"
    payload = json.dumps(
        {
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "think": False,
            "stream": False,
            "options": {"num_ctx": 32768},
        }
    ).encode("utf-8")
    req = urllib.request.Request(
        url, data=payload, headers={"Content-Type": "application/json"}, method="POST"
    )
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            data = json.loads(resp.read())
        return data["message"]["content"]
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="replace")
        return f"FEHLER: HTTP {e.code}: {body[:400]}"
    except Exception as e:
        if cancel_event is not None and cancel_event.is_set():
            return "FEHLER: Abgebrochen"
        return f"FEHLER: Ollama-Aufruf fehlgeschlagen: {e}"


def _call_openai_compat(
    base_url: str,
    api_key: str,
    model: str,
    prompt: str,
    timeout: int,
    extra_body: dict | None = None,
    cancel_event: threading.Event | None = None,
) -> str:
    """Ruft eine OpenAI-kompatible Chat-API auf (GLM, Kimi, OpenAI, Ollama)."""
    if cancel_event is not None and cancel_event.is_set():
        return "FEHLER: Abgebrochen"
    url = base_url.rstrip("/") + "/chat/completions"
    body: dict = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 4096,
    }
    if extra_body:
        body.update(extra_body)
    payload = json.dumps(body).encode("utf-8")
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    req = urllib.request.Request(url, data=payload, headers=headers, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            data = json.loads(resp.read())
        return data["choices"][0]["message"]["content"]
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="replace")
        return f"FEHLER: HTTP {e.code}: {body[:400]}"
    except Exception as e:
        if cancel_event is not None and cancel_event.is_set():
            return "FEHLER: Abgebrochen"
        return f"FEHLER: API-Aufruf fehlgeschlagen: {e}"


def api_key_available(config: dict[str, Any]) -> bool:
    """Prüft ob der API-Key für den konfigurierten Provider gesetzt ist."""
    provider = config.get("api", {}).get("provider", "anthropic").lower()
    key_map = {
        "anthropic": "ANTHROPIC_API_KEY",
        "glm": "GLM_API_KEY",
        "kimi": "KIMI_API_KEY",
        "openai": "OPENAI_API_KEY",
        "ollama": None,  # kein Key nötig
    }
    env_key = key_map.get(provider)
    if env_key is None:
        return True
    return bool(os.environ.get(env_key, ""))


def run_llm_api(
    prompt: str,
    config: dict[str, Any],
    cancel_event: threading.Event | None = None,
    schema: dict[str, Any] | None = None,
) -> str:
    """Dispatcht den LLM-Aufruf je nach konfiguriertem Provider."""
    if cancel_event is not None and cancel_event.is_set():
        return "FEHLER: Abgebrochen"
    api_cfg = config.get("api", {})
    provider = api_cfg.get("provider", "anthropic").lower()
    timeout = config.get("agent", {}).get("timeout_seconds", 120)
    model = api_cfg.get("model", "")

    if provider == "anthropic":
        return run_anthropic_api(
            prompt, model or "claude-sonnet-4-6", timeout, cancel_event=cancel_event,
            schema=schema,
        )

    if provider == "glm":
        api_key = os.environ.get("GLM_API_KEY", "")
        if not api_key:
            return "FEHLER: GLM_API_KEY nicht gesetzt (.env prüfen)"
        base_url = os.environ.get("GLM_BASE_URL", "https://open.bigmodel.cn/api/paas/v4")
        return _call_openai_compat(
            base_url,
            api_key,
            model or "glm-4-flash",
            prompt,
            timeout,
            cancel_event=cancel_event,
        )

    if provider == "kimi":
        api_key = os.environ.get("KIMI_API_KEY", "")
        if not api_key:
            return "FEHLER: KIMI_API_KEY nicht gesetzt (.env prüfen)"
        base_url = os.environ.get("KIMI_BASE_URL", "https://api.moonshot.ai/v1")
        return _call_openai_compat(
            base_url,
            api_key,
            model or "moonshot-v1-8k",
            prompt,
            timeout,
            cancel_event=cancel_event,
        )

    if provider == "openai":
        api_key = os.environ.get("OPENAI_API_KEY", "")
        if not api_key:
            return "FEHLER: OPENAI_API_KEY nicht gesetzt (.env prüfen)"
        return _call_openai_compat(
            "https://api.openai.com/v1",
            api_key,
            model or "gpt-4o-mini",
            prompt,
            timeout,
            cancel_event=cancel_event,
        )

    if provider == "ollama":
        base_url = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434")
        return _call_ollama_native(
            base_url,
            model or "qwen3.5:27b",
            prompt,
            timeout,
            cancel_event=cancel_event,
        )

    return (
        f"FEHLER: Unbekannter Provider '{provider}' — erlaubt: anthropic, glm, kimi, openai, ollama"
    )


def run_anthropic_api(
    prompt: str,
    model: str,
    timeout: int = 120,
    cancel_event: threading.Event | None = None,
    schema: dict[str, Any] | None = None,
) -> str:
    if cancel_event is not None and cancel_event.is_set():
        return "FEHLER: Abgebrochen"
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return "FEHLER: ANTHROPIC_API_KEY nicht gesetzt"
    try:
        import anthropic
    except ImportError:
        return "FEHLER: anthropic-Paket nicht installiert (pip install anthropic)"
    client = anthropic.Anthropic(api_key=api_key)
    try:
        # Tool Use (Structured Output): garantiert schema-konformes JSON
        if schema:
            response = client.messages.create(
                model=model,
                max_tokens=4096,
                tools=[{
                    "name": "feedback_result",
                    "description": "Bewertungsergebnis der Schülerarbeit als strukturiertes JSON",
                    "input_schema": schema,
                }],
                tool_choice={"type": "tool", "name": "feedback_result"},
                messages=[{"role": "user", "content": prompt}],
                timeout=timeout,
            )
            for block in response.content:
                if block.type == "tool_use":
                    return json.dumps(block.input, ensure_ascii=False)
            return "FEHLER: Kein Tool-Use-Block in Antwort"
        # Fallback: normaler Text-Aufruf (für nicht-Schema-Prompts)
        response = client.messages.create(
            model=model,
            max_tokens=4096,
            messages=[{"role": "user", "content": prompt}],
            timeout=timeout,
        )
        if response.content and len(response.content) > 0:
            return response.content[0].text
        return "FEHLER: Leere Antwort"
    except Exception as e:
        if cancel_event is not None and cancel_event.is_set():
            return "FEHLER: Abgebrochen"
        return f"FEHLER: API-Aufruf fehlgeschlagen: {e}"


def check_agent_availability(config: dict[str, Any]) -> dict[str, bool]:
    commands = config.get("agent", {}).get("commands", {})
    availability: dict[str, bool] = {}
    for name, cmd in commands.items():
        binary = cmd.split()[0]
        try:
            result = subprocess.run(
                ["which", binary],
                capture_output=True,
                timeout=2,
            )
            availability[name] = result.returncode == 0
        except (FileNotFoundError, subprocess.TimeoutExpired):
            availability[name] = False
    return availability


def humanize_agent_error(error: str, agent_name: str) -> str:
    if "nicht gefunden" in error or "not found" in error.lower():
        return (
            f"Der Agent {agent_name!r} ist nicht installiert.\n"
            f"  Tipp: Nutze einen anderen Agent oder den Zwischenablage-Modus."
        )
    if "Timeout" in error or "timeout" in error.lower():
        return (
            f"Der Agent {agent_name!r} hat zu lange gebraucht.\n"
            f"  Tipp: Erhoehe den Timeout in den Einstellungen oder nutze einen anderen Agent."
        )
    if "API" in error and "key" in error.lower():
        return (
            f"API-Schluessel fehlt oder ist ungueltig.\n"
            f"  Tipp: Setze ANTHROPIC_API_KEY und pruefe die Verbindung im Einstellungsmenue."
        )
    return error


def compute_statistics(analyses: list[dict[str, Any]]) -> dict[str, Any]:
    """Berechnet Klassen-Statistiken aus einer Liste von Analyse-Dicts.

    Returns:
        {
            "total": int,                         # Anzahl ausgewerteter Dateien
            "grade_distribution": {1..5: int},    # Anzahl pro Note
            "grade_average": float,               # Gewichteter Gesamtdurchschnitt
            "criteria_averages": {                # Pro Kriterium
                key: {"avg": float, "count": int, "min": float, "max": float}
            },
            "weakest_criterion": str | None,      # Kriterium mit niedrigstem Ø
            "strongest_criterion": str | None,    # Kriterium mit höchstem Ø
        }
    """
    grade_dist: dict[int, int] = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
    criteria_scores: dict[str, list[float]] = {}

    for data in analyses:
        note_data = data.get("notenempfehlung", {})
        note = note_data.get("note")
        if isinstance(note, (int, float)) and 1 <= int(note) <= 5:
            grade_dist[int(note)] += 1

        bewertung = data.get("bewertung", {})
        for key, crit in bewertung.items():
            if isinstance(crit, dict):
                punkte = crit.get("punkte")
                if isinstance(punkte, (int, float)):
                    criteria_scores.setdefault(key, []).append(float(punkte))

    total = sum(grade_dist.values())
    grade_average = (
        round(sum(n * c for n, c in grade_dist.items()) / total, 2) if total > 0 else 0.0
    )

    criteria_averages = {
        key: {
            "avg": round(sum(scores) / len(scores), 2),
            "count": len(scores),
            "min": min(scores),
            "max": max(scores),
        }
        for key, scores in criteria_scores.items()
    }

    weakest = (
        min(criteria_averages, key=lambda k: criteria_averages[k]["avg"])
        if criteria_averages
        else None
    )
    strongest = (
        max(criteria_averages, key=lambda k: criteria_averages[k]["avg"])
        if criteria_averages
        else None
    )

    return {
        "total": total,
        "grade_distribution": grade_dist,
        "grade_average": grade_average,
        "criteria_averages": criteria_averages,
        "weakest_criterion": weakest,
        "strongest_criterion": strongest,
    }


def compute_class_progress(config: dict[str, Any], klasse: str) -> list[dict[str, Any]]:
    """Berechnet Lernfortschritt ueber alle Aufgaben einer Klasse hinweg.

    Returns:
        Liste von Dicts (eins pro Aufgabe mit Analysen), sortiert nach
        Config-Reihenfolge:
            [{"aufgabe": slug, "label": ..., "avg_note": float,
              "avg_criteria": {key: float}, "n": int}, ...]
    """
    aufgaben = list_aufgaben(config, klasse)
    progress: list[dict[str, Any]] = []

    for slug in aufgaben:
        auf_cfg = get_aufgabe_cfg(config, klasse, slug)
        label = auf_cfg.get("label", slug)
        paths = build_project_paths(config, klasse, slug)
        fb_dir = paths.feedback_data_dir

        if not fb_dir.exists():
            continue

        analyses: list[dict[str, Any]] = []
        for json_path in sorted(fb_dir.glob("*.json")):
            try:
                data = json.loads(json_path.read_text(encoding="utf-8"))
                if isinstance(data, dict) and "notenempfehlung" in data:
                    analyses.append(data)
            except (json.JSONDecodeError, OSError):
                continue

        if not analyses:
            continue

        stats = compute_statistics(analyses)

        criteria_avgs = {key: vals["avg"] for key, vals in stats["criteria_averages"].items()}

        progress.append(
            {
                "aufgabe": slug,
                "label": label,
                "avg_note": stats["grade_average"],
                "avg_criteria": criteria_avgs,
                "n": stats["total"],
            }
        )

    return progress


def docx_to_pdf(docx_path: Path, out_path: Path | None = None) -> Path | None:
    """Konvertiert DOCX nach PDF via LibreOffice headless oder unoconv.

    Args:
        docx_path: Pfad zur Quelldatei (.docx)
        out_path: Gewuenschter Ausgabepfad (optional; Standard: gleicher Ordner, .pdf-Endung)

    Returns:
        Pfad zur erzeugten PDF-Datei, oder None wenn kein Konverter verfuegbar.
    """
    candidates = [
        (shutil.which("libreoffice"), ["libreoffice", "--headless", "--convert-to", "pdf"]),
        (shutil.which("soffice"), ["soffice", "--headless", "--convert-to", "pdf"]),
        (shutil.which("unoconv"), ["unoconv", "-f", "pdf"]),
    ]

    cmd_base = None
    for found, base in candidates:
        if found:
            cmd_base = base
            break

    if cmd_base is None:
        return None

    if out_path is None:
        out_path = docx_path.with_suffix(".pdf")

    out_dir = str(out_path.parent)

    if cmd_base[0] in ("libreoffice", "soffice"):
        cmd = cmd_base + [f"--outdir={out_dir}", str(docx_path)]
    else:
        cmd = cmd_base + ["-o", str(out_path), str(docx_path)]

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if result.returncode == 0 and out_path.exists():
            return out_path
    except (subprocess.TimeoutExpired, FileNotFoundError, OSError):
        pass

    return None
