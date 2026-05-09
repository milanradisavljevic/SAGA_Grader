from __future__ import annotations

import json
from argparse import Namespace
from pathlib import Path
import sys

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import generate_feedback as gf


FIXTURES = Path(__file__).parent / "fixtures"


def make_paths(tmp_path: Path) -> gf.ProjectPaths:
    output_dir = tmp_path / "output"
    return gf.ProjectPaths(
        root=tmp_path,
        input_dir=tmp_path / "input",
        output_dir=output_dir,
        feedback_data_dir=output_dir / "feedback_data",
        fehlerlog=output_dir / "fehlerlog.txt",
    )


def test_load_feedback_json_parses_fixture() -> None:
    data = gf.load_feedback_json(FIXTURES / "tamara_feedback.json")

    assert data.schueler == "Anna Muster"
    assert data.fach == "Deutsch"
    assert len(data.bewertung) == 4
    assert data.notenempfehlung.note == 3


def test_build_feedback_document_contains_expected_sections() -> None:
    data = gf.load_feedback_json(FIXTURES / "emma_b2_feedback.json")
    document = gf.build_feedback_document(data)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "KORREKTUR-FEEDBACK" in text
    assert "TASK ACHIEVEMENT" in text
    assert "Empfohlene Note: 2 - Gut" in text


def test_process_file_creates_docx_and_can_be_reopened(tmp_path: Path) -> None:
    paths = make_paths(tmp_path)
    paths.feedback_data_dir.mkdir(parents=True, exist_ok=True)
    paths.output_dir.mkdir(parents=True, exist_ok=True)

    source = FIXTURES / "matthias_feedback.json"
    target = paths.feedback_data_dir / source.name
    target.write_text(source.read_text(encoding="utf-8"), encoding="utf-8")

    message = gf.process_file(target, paths, force=False, dry_run=False)

    output_path = (
        paths.output_dir / "Kommentar -Krass, Digga!...- Max Mustermann_feedback.docx"
    )
    assert "Gespeichert" in message
    assert output_path.exists()

    reopened = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
    assert "STIL UND AUSDRUCK" in text
    assert "Rhetorische Figuren:" in text
    assert "Fehler im Detail:" in text


def test_process_file_skips_existing_output_without_force(tmp_path: Path) -> None:
    paths = make_paths(tmp_path)
    paths.feedback_data_dir.mkdir(parents=True, exist_ok=True)
    paths.output_dir.mkdir(parents=True, exist_ok=True)

    source = FIXTURES / "tamara_feedback.json"
    target = paths.feedback_data_dir / source.name
    target.write_text(source.read_text(encoding="utf-8"), encoding="utf-8")
    existing = paths.output_dir / "deutsch digga_feedback.docx"
    existing.write_text("placeholder", encoding="utf-8")

    message = gf.process_file(target, paths, force=False, dry_run=False)

    assert "Uebersprungen" in message
    assert existing.read_text(encoding="utf-8") == "placeholder"


def test_main_logs_invalid_json(tmp_path: Path, monkeypatch) -> None:
    paths = make_paths(tmp_path)
    paths.feedback_data_dir.mkdir(parents=True, exist_ok=True)
    paths.output_dir.mkdir(parents=True, exist_ok=True)

    invalid = paths.feedback_data_dir / "broken.json"
    invalid.write_text(json.dumps({"datei": "broken.docx"}), encoding="utf-8")

    monkeypatch.setattr(gf, "project_paths", lambda: paths)
    monkeypatch.setattr(
        gf, "parse_args", lambda: Namespace(file=None, force=False, dry_run=False)
    )

    exit_code = gf.main()

    assert exit_code == 1
    assert "broken.json - Fehler" in paths.fehlerlog.read_text(encoding="utf-8")


def test_summary_table_present_in_output() -> None:
    data = gf.load_feedback_json(FIXTURES / "tamara_feedback.json")
    doc = gf.build_feedback_document(data)
    assert len(doc.tables) >= 2
    last_table = doc.tables[-1]
    last_row_text = last_table.rows[-1].cells[0].text
    assert "GESAMTNOTE" in last_row_text


def test_grade_color_five_stages() -> None:
    assert gf._note_color(1) == gf.C_GRADE_1  # hellgrün
    assert gf._note_color(2) == gf.C_GRADE_2  # dunkelgrün
    assert gf._note_color(3) == gf.C_GRADE_3  # gelb
    assert gf._note_color(4) == gf.C_GRADE_4  # orange
    assert gf._note_color(5) == gf.C_GRADE_5  # rot
    # Rückwärtskompatible Aliasse
    assert gf.C_GRADE_GOOD == gf.C_GRADE_2
    assert gf.C_GRADE_OK == gf.C_GRADE_3
    assert gf.C_GRADE_FAIL == gf.C_GRADE_5


def test_page_format_is_a4() -> None:
    data = gf.load_feedback_json(FIXTURES / "tamara_feedback.json")
    doc = gf.build_feedback_document(data)
    section = doc.sections[0]
    assert abs(section.page_width.cm - 21.0) < 0.5
    assert abs(section.page_height.cm - 29.7) < 0.5


def test_footer_contains_saga() -> None:
    data = gf.load_feedback_json(FIXTURES / "tamara_feedback.json")
    doc = gf.build_feedback_document(data)
    footer = doc.sections[0].footer
    text = " ".join(p.text for p in footer.paragraphs)
    assert "SAGA" in text


def test_header_contains_datei_without_config() -> None:
    data = gf.load_feedback_json(FIXTURES / "tamara_feedback.json")
    doc = gf.build_feedback_document(data, config=None)
    all_table_text = " ".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    assert "deutsch digga.docx" in all_table_text


def test_document_header_with_config() -> None:
    data = gf.load_feedback_json(FIXTURES / "tamara_feedback.json")
    config = {
        "docx": {
            "teacher_name": "Mag. Mueller",
            "school_name": "BRG Wien",
            "logo_path": "",
        }
    }
    doc = gf.build_feedback_document(data, config=config)
    header = doc.sections[0].header
    header_text = " ".join(p.text for p in header.paragraphs)
    assert "Mag. Mueller" in header_text
    assert "BRG Wien" in header_text
